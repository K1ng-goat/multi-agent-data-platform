import io
import os
import re
from pathlib import Path

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, NamedStyle
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

import style_config as sc
import theme_service as ts

EXPORTS_DIR = Path(__file__).resolve().parent / "exports"

# ── CJK font detection for matplotlib ─────────────────
_CJK_FONT = None
for _fname in fm.findSystemFonts():
    try:
        _prop = fm.FontProperties(fname=_fname)
        _name = _prop.get_name()
        if any(kw in _name.lower() for kw in ["microsoft yahei", "simhei", "simsun", "noto sans cjk", "wenquanyi"]):
            _CJK_FONT = _prop
            break
    except Exception:
        pass
if _CJK_FONT is None:
    try:
        _CJK_FONT = fm.FontProperties(family="sans-serif")
    except Exception:
        _CJK_FONT = fm.FontProperties()

# ── Matplotlib RC defaults ─────────────────────────────
plt.rcParams["axes.grid"] = True
plt.rcParams["grid.alpha"] = sc.CHART_GRID_ALPHA
plt.rcParams["axes.edgecolor"] = "#cccccc"
plt.rcParams["axes.titlesize"] = sc.CHART_TITLE_SIZE
plt.rcParams["axes.titleweight"] = "bold"
plt.rcParams["figure.dpi"] = sc.CHART_DPI


def _ensure_dir():
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)


def _apply_style(session: dict):
    """Apply the session's theme + overrides to style_config globals."""
    cfg = ts.get_effective_config(session)
    sc.load_from_dict(cfg)


def _base_name(session: dict) -> str:
    """Extract base filename without extension, e.g. 'sales.xlsx' → 'sales'."""
    ds = session.get("data_summary", {})
    filename = ds.get("filename", "output")
    base = Path(filename).stem
    return base or "output"


def _sanitize_filename(s: str) -> str:
    """Remove characters unsafe for filenames."""
    return re.sub(r'[\\/*?:"<>|]', "_", s)


def _thin_border() -> Border:
    return Border(
        left=Side(style="thin", color=sc.BORDER_COLOR),
        right=Side(style="thin", color=sc.BORDER_COLOR),
        top=Side(style="thin", color=sc.BORDER_COLOR),
        bottom=Side(style="thin", color=sc.BORDER_COLOR),
    )


# ═══════════════════════════════════════════════════════
#  Excel Export
# ═══════════════════════════════════════════════════════

def export_excel(session: dict) -> str:
    """Export analysis results to a formatted Excel file."""
    _ensure_dir()
    _apply_style(session)
    df: pd.DataFrame = session["df"]
    analysis = session.get("analysis", {})
    base_name = _base_name(session)
    path = str(EXPORTS_DIR / f"{_sanitize_filename(base_name)}_统计结果.xlsx")
    wb = Workbook()

    header_font = Font(
        name=sc.EXCEL_HEADER_FONT_NAME, bold=True,
        size=sc.EXCEL_HEADER_FONT_SIZE, color=sc.HEADER_FG_COLOR,
    )
    header_fill = PatternFill(start_color=sc.HEADER_BG_COLOR, end_color=sc.HEADER_BG_COLOR, fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    body_font = Font(name=sc.EXCEL_BODY_FONT_NAME, size=sc.EXCEL_BODY_FONT_SIZE)
    body_align = Alignment(horizontal="center", vertical="center")

    def _style_header(ws, row, ncols):
        for c in range(1, ncols + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = _thin_border()

    def _auto_width(ws, ncols, data_rows, min_w=None, max_w=None):
        min_w = min_w or sc.EXCEL_MIN_COL_WIDTH
        max_w = max_w or sc.EXCEL_MAX_COL_WIDTH
        for c in range(1, ncols + 1):
            col_letter = get_column_letter(c)
            max_len = 0
            for row in range(1, data_rows + 1):
                val = ws.cell(row=row, column=c).value
                if val:
                    # Rough CJK width: each CJK char ≈ 2 units
                    width = sum(2 if ord(ch) > 127 else 1 for ch in str(val))
                    max_len = max(max_len, width)
            ws.column_dimensions[col_letter].width = max(min_w, min(max_len + 3, max_w))

    # ── Sheet 1: Original data ───────────────────────
    ws1 = wb.active
    ws1.title = "原始数据"
    ncols = len(df.columns)
    for c, col_name in enumerate(df.columns, 1):
        ws1.cell(row=1, column=c, value=str(col_name))
    for r, row in enumerate(df.fillna("").values.tolist(), 2):
        for c, val in enumerate(row, 1):
            cell = ws1.cell(row=r, column=c, value=val)
            cell.font = body_font
            cell.alignment = body_align
            cell.border = _thin_border()
    _style_header(ws1, 1, ncols)
    _auto_width(ws1, ncols, len(df) + 1)
    ws1.freeze_panes = f"A{sc.EXCEL_FREEZE_ROW + 1}"
    ws1.auto_filter.ref = f"A1:{get_column_letter(ncols)}{len(df) + 1}"

    # ── Sheet 2: AI Analysis ─────────────────────────
    ws2 = wb.create_sheet("AI分析")
    section_font = Font(name=sc.EXCEL_HEADER_FONT_NAME, bold=True, size=13, color=sc.PRIMARY_COLOR)
    content_font = Font(name=sc.EXCEL_BODY_FONT_NAME, size=11)
    sections = [
        ("数据总结", analysis.get("summary", "暂无")),
        ("异常分析", analysis.get("anomaly", "暂无")),
        ("趋势分析", analysis.get("trend", "暂无")),
    ]
    row = 1
    for title, content in sections:
        ws2.cell(row=row, column=1, value=title).font = section_font
        row += 1
        ws2.merge_cells(start_row=row, start_column=1, end_row=row + 1, end_column=6)
        cell = ws2.cell(row=row, column=1, value=content)
        cell.font = content_font
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws2.row_dimensions[row].height = 60
        row += 3
    ws2.column_dimensions["A"].width = 20
    for c in range(2, 7):
        ws2.column_dimensions[get_column_letter(c)].width = 18

    # ── Sheet 3: Statistics ──────────────────────────
    ws3 = wb.create_sheet("统计信息")
    numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
    if numeric_cols:
        stats = df[numeric_cols].describe()
        ncols3 = len(stats.columns) + 1
        ws3.cell(row=1, column=1, value="统计项")
        for c, col_name in enumerate(stats.columns, 2):
            ws3.cell(row=1, column=c, value=str(col_name))
        _style_header(ws3, 1, ncols3)
        for r, (idx, row_data) in enumerate(stats.iterrows(), 2):
            cell = ws3.cell(row=r, column=1, value=str(idx))
            cell.font = Font(name=sc.EXCEL_BODY_FONT_NAME, bold=True, size=sc.EXCEL_BODY_FONT_SIZE)
            cell.border = _thin_border()
            cell.alignment = body_align
            for c, val in enumerate(row_data.tolist(), 2):
                cell = ws3.cell(row=r, column=c, value=round(float(val), 2) if not pd.isna(val) else "")
                cell.font = body_font
                cell.border = _thin_border()
                cell.alignment = body_align
        _auto_width(ws3, ncols3, len(stats) + 1)
        ws3.freeze_panes = f"A{sc.EXCEL_FREEZE_ROW + 1}"

    wb.save(path)
    return path


# ═══════════════════════════════════════════════════════
#  Word Export
# ═══════════════════════════════════════════════════════

def _set_cjk_font(run, font_name, font_size, bold=False):
    """Set both Western and East-Asian font on a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _set_paragraph_spacing(para, line_spacing=None, after=None, before=None):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if after is not None:
        pf.space_after = after
    if before is not None:
        pf.space_before = before


def export_word(session: dict) -> str:
    """Generate a professionally formatted Word report."""
    _ensure_dir()
    _apply_style(session)
    ds = session.get("data_summary", {})
    analysis = session.get("analysis", {})
    charts = session.get("charts", [])
    base_name = _base_name(session)
    path = str(EXPORTS_DIR / f"{_sanitize_filename(base_name)}_分析报告.docx")
    doc = Document()

    # ── Page margins ──────────────────────────────────
    for section in doc.sections:
        section.top_margin = sc.PAGE_MARGIN_TOP
        section.bottom_margin = sc.PAGE_MARGIN_BOTTOM
        section.left_margin = sc.PAGE_MARGIN_LEFT
        section.right_margin = sc.PAGE_MARGIN_RIGHT

    # ── Default style ─────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = sc.BODY_FONT_NAME
    style.font.size = sc.BODY_SIZE
    style.paragraph_format.line_spacing = sc.LINE_SPACING
    style.paragraph_format.space_after = sc.PARAGRAPH_SPACING_AFTER
    # Set East-Asian font on Normal style
    rPr = style.element.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            rFonts.set(qn("w:eastAsia"), sc.BODY_FONT_NAME)

    # ── Title ─────────────────────────────────────────
    title = doc.add_heading("数据分析报告", level=0)
    title.alignment = sc.TITLE_ALIGNMENT
    _set_paragraph_spacing(title, after=Pt(16))
    for run in title.runs:
        _set_cjk_font(run, sc.TITLE_FONT_NAME, sc.TITLE_SIZE, bold=True)

    # ── Section 1: Basic Info ─────────────────────────
    h1 = doc.add_heading("一、基本信息", level=1)
    _set_paragraph_spacing(h1, after=Pt(8), before=Pt(12))
    for run in h1.runs:
        _set_cjk_font(run, sc.TITLE_FONT_NAME, sc.H1_SIZE, bold=True)

    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = doc.styles["Table Grid"]
    info_data = [
        ("文件名", ds.get("filename", "未知")),
        ("数据规模", f"{ds.get('shape', {}).get('rows', 0)} 行 x {ds.get('shape', {}).get('columns', 0)} 列"),
        ("数据列", ", ".join(ds.get("columns", []))),
    ]
    for i, (key, val) in enumerate(info_data):
        for j, text in enumerate([key, str(val)]):
            cell = info_table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            _set_cjk_font(run, sc.TABLE_FONT_NAME, sc.TABLE_SIZE, bold=(j == 0))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_spacing(p, line_spacing=1.2, after=Pt(2))

    # ── Helper to add a section ──────────────────────
    def _add_section(heading_text: str, body_text: str):
        h = doc.add_heading(heading_text, level=1)
        _set_paragraph_spacing(h, after=Pt(8), before=Pt(12))
        for run in h.runs:
            _set_cjk_font(run, sc.TITLE_FONT_NAME, sc.H1_SIZE, bold=True)
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, line_spacing=sc.LINE_SPACING, after=Pt(8))
        run = p.add_run(body_text or "暂无")
        _set_cjk_font(run, sc.BODY_FONT_NAME, sc.BODY_SIZE, bold=False)
        return p

    _add_section("二、数据总结", analysis.get("summary", "暂无数据总结"))
    _add_section("三、异常分析", analysis.get("anomaly", "暂无异常数据"))
    _add_section("四、趋势分析", analysis.get("trend", "暂无趋势分析"))
    _add_section("五、关键发现与建议", "基于以上数据分析，请参考 AI 分析结果中的关键指标和趋势说明。")

    # ── Charts ────────────────────────────────────────
    if charts:
        doc.add_page_break()
        h = doc.add_heading("六、数据图表", level=1)
        _set_paragraph_spacing(h, after=Pt(12), before=Pt(8))
        for run in h.runs:
            _set_cjk_font(run, sc.TITLE_FONT_NAME, sc.H1_SIZE, bold=True)

        for i, chart in enumerate(charts):
            chart_title = chart.get("title", f"图表{i + 1}")
            safe_title = _sanitize_filename(chart_title)
            chart_path = _save_chart_image(chart, f"{base_name}_{safe_title}")

            h2 = doc.add_heading(chart_title, level=2)
            _set_paragraph_spacing(h2, after=Pt(6), before=Pt(8))
            for run in h2.runs:
                _set_cjk_font(run, sc.TITLE_FONT_NAME, sc.H2_SIZE, bold=True)

            if chart_path and os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5.8))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_spacing(last_para, after=Pt(4))
            else:
                p = doc.add_paragraph("(图表生成失败)")
                _set_paragraph_spacing(p, after=Pt(4))

    doc.save(path)
    return path


# ═══════════════════════════════════════════════════════
#  Chart Export
# ═══════════════════════════════════════════════════════

def _save_chart_image(chart: dict, filename: str) -> str | None:
    """Render a chart config to a PNG file using matplotlib."""
    _ensure_dir()
    chart_type = chart.get("type", "line")
    title = chart.get("title", "")
    labels = chart.get("labels", [])
    datasets = chart.get("datasets", [])

    if not labels or not datasets:
        return None

    fig, ax = plt.subplots(figsize=sc.CHART_FIGSIZE)
    colors = sc.CHART_COLORS

    if chart_type == "pie":
        data_values = datasets[0].get("data", []) if datasets else []
        wedges, texts, autotexts = ax.pie(
            data_values, labels=labels, autopct="%1.1f%%",
            colors=colors[:len(labels)], startangle=90,
            textprops={"fontsize": sc.CHART_LABEL_SIZE},
        )
        if _CJK_FONT:
            for t in texts + autotexts:
                t.set_fontproperties(_CJK_FONT)
        # White border on wedges for separation
        for w in wedges:
            w.set_edgecolor("white")
            w.set_linewidth(1.2)
    else:
        x = range(len(labels))
        n_datasets = max(len(datasets), 1)
        width = 0.75 / n_datasets
        for j, ds in enumerate(datasets):
            data_values = ds.get("data", [])
            offset = (j - (n_datasets - 1) / 2) * width
            positions = [i + offset for i in x]
            if chart_type == "bar":
                ax.bar(positions, data_values, width=width * 0.88,
                       label=ds.get("name", ""), color=colors[j % len(colors)],
                       edgecolor="white", linewidth=0.5)
            else:
                ax.plot(positions, data_values, marker="o", linewidth=sc.CHART_LINE_WIDTH,
                        markersize=sc.CHART_MARKER_SIZE,
                        label=ds.get("name", ""), color=colors[j % len(colors)])
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=40, ha="right", fontsize=sc.CHART_LABEL_SIZE)
        if _CJK_FONT:
            for lbl in ax.get_xticklabels():
                lbl.set_fontproperties(_CJK_FONT)
        legend = ax.legend(loc="upper right", fontsize=sc.CHART_LEGEND_SIZE,
                          frameon=True, fancybox=True, framealpha=0.9,
                          edgecolor="#dddddd")
        if _CJK_FONT and legend:
            for txt in legend.get_texts():
                txt.set_fontproperties(_CJK_FONT)
        # Remove top/right spines
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

    ax.set_title(title, fontsize=sc.CHART_TITLE_SIZE, fontweight="bold", pad=12)
    if _CJK_FONT:
        ax.title.set_fontproperties(_CJK_FONT)
    ax.set_xlabel("")
    ax.set_ylabel("")
    ax.grid(True, alpha=sc.CHART_GRID_ALPHA, linestyle="--")

    fig.tight_layout()
    path = str(EXPORTS_DIR / f"{_sanitize_filename(filename)}.png")
    fig.savefig(path, dpi=sc.CHART_DPI, bbox_inches="tight")
    plt.close(fig)
    return path


def export_chart(chart: dict, index: int, base_name: str, fmt: str = "png", session: dict | None = None) -> str | None:
    """Export a single chart as an image file."""
    _ensure_dir()
    if session:
        _apply_style(session)
    ext = fmt if fmt in ("png", "jpg") else "png"
    chart_title = chart.get("title", f"图表{index + 1}")
    safe_title = _sanitize_filename(chart_title)
    fname_base = f"{_sanitize_filename(base_name)}_{safe_title}"
    path = str(EXPORTS_DIR / f"{fname_base}.{ext}")
    fig_path = _save_chart_image(chart, fname_base)
    if fig_path is None:
        return None
    if ext == "jpg":
        from PIL import Image
        img = Image.open(fig_path)
        img = img.convert("RGB")
        img.save(path, "JPEG", quality=90)
        return path
    return fig_path
